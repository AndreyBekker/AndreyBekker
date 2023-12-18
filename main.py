import docx

for i in range(1, 6):
    doc = docx.Document(f'{i}.docx')

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = docx.shared.Pt(14)

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.5

    doc.save(f'{i}_updated.docx')
